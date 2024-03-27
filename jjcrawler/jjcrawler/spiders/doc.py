from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .utils import get_file_name, get_heading, format_body


def set_doc_style(doc):
    section = doc.sections[0]
    section.left_margin, section.right_margin = Cm(2), Cm(2)

    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Helvetica"
    normal_font.name = "Microsoft YaHei"
    normal_font.size = Pt(12)
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.line_spacing = Pt(14)


def create_desc_doc(directory, novel):
    print(f"下载 文案 中...")
    description_doc = Document()
    file_name = "文案"
    set_doc_style(description_doc)

    for paragraph in novel["desc"]:
        description_doc.add_paragraph(paragraph)

    tags = "内容标签： " + " ".join(novel["tag_list"])
    tags_paragraph = description_doc.add_paragraph(tags)
    tags_paragraph.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    keywords_paragraph = description_doc.add_paragraph(novel["keywords"])
    keywords_paragraph.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    oneliner_paragraph = description_doc.add_paragraph(novel["oneliner"])
    oneliner_paragraph.runs[0].font.color.rgb = RGBColor(0xF9, 0x8C, 0x4D)

    meaning_paragraph = description_doc.add_paragraph(novel["meaning"])
    meaning_paragraph.runs[0].font.color.rgb = RGBColor(0xF9, 0x8C, 0x4D)

    output_path = f"{directory}{file_name}.docx"
    description_doc.save(output_path)


def create_chapter_doc(directory, chapter):
    chapter_doc = Document()
    file_name = get_file_name(chapter)
    print(f"下载 {file_name} 中...")

    set_doc_style(chapter_doc)

    heading = get_heading(chapter)
    heading_paragraph = chapter_doc.add_paragraph(heading)
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    chapter_doc.add_paragraph()

    for paragraph in format_body(chapter["body"]):
        chapter_doc.add_paragraph(paragraph)

    author_said = chapter["author_said"]
    if author_said:
        chapter_doc.add_paragraph()
        for author_said_p in author_said:
            author_said_paragraph = chapter_doc.add_paragraph(author_said_p)
            if author_said_paragraph.runs:
                author_said_paragraph.runs[0].font.color.rgb = RGBColor(
                    0x00, 0x99, 0x00
                )
    output_path = f"{directory}{file_name}.docx"
    chapter_doc.save(output_path)
